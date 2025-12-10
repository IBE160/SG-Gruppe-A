import logging
from docx import Document
import os

logger = logging.getLogger(__name__)

def parse_cv(file_path: str) -> str:
    """
    Parses a CV file and extracts text content.
    Currently supports .docx files using python-docx.
    
    Args:
        file_path: Path to the file to parse
        
    Returns:
        Extracted text content or empty string if parsing fails/unsupported format.
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found at {file_path}")
        return ""

    # python-docx only supports .docx
    # We check extension case-insensitively
    if not file_path.lower().endswith(".docx"):
        logger.warning(f"Skipping parsing for non-docx file: {file_path}")
        return ""

    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        
        # Also extract text from tables if necessary? 
        # CVs often use tables for layout.
        # python-docx paragraphs might not include table content?
        # Paragraphs in Document.paragraphs are body paragraphs.
        # Tables are in Document.tables.
        
        # Let's add table text as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            full_text.append(text)
                            
        return "\n\n".join(full_text)
        
    except Exception as e:
        logger.error(f"Error parsing .docx file {file_path}: {e}", exc_info=True)
        return ""
