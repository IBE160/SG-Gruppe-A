from docx import Document

def create_mock_cv():
    document = Document()
    document.add_heading('John Doe', 0)
    document.add_paragraph('Software Developer')
    document.add_heading('Experience', level=1)
    document.add_paragraph('Senior Developer at Tech Corp (2020-Present)')
    document.add_paragraph('Developed web applications using Python and React.')
    document.add_heading('Education', level=1)
    document.add_paragraph('B.Sc. Computer Science, University of Technology')
    document.add_heading('Skills', level=1)
    document.add_paragraph('Python, FastAPI, React, SQL, AI integration')
    
    document.save('poc/sample_cv.docx')
    print("Created poc/sample_cv.docx")

if __name__ == "__main__":
    create_mock_cv()
