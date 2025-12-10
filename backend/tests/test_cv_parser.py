import pytest
from app.services.cv_parser import parse_cv
from docx import Document
import os

def test_parse_docx_success(tmp_path):
    # Create a dummy docx
    d = tmp_path / "test_cv.docx"
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.save(d)
    
    text = parse_cv(str(d))
    assert "First paragraph" in text
    assert "Second paragraph" in text

def test_parse_docx_with_table(tmp_path):
    d = tmp_path / "test_cv_table.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Table content"
    doc.save(d)
    
    text = parse_cv(str(d))
    # Note: Our implementation adds table text after paragraphs
    assert "Table content" in text

def test_parse_non_existent_file():
    text = parse_cv("non_existent.docx")
    assert text == ""

def test_parse_non_docx_extension(tmp_path):
    d = tmp_path / "test.txt"
    d.write_text("hello", encoding="utf-8")
    text = parse_cv(str(d))
    assert text == ""
